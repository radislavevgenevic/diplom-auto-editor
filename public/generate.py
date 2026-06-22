import json
import os
import shutil
import tempfile
from zipfile import ZipFile
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE

PLACEHOLDERS = {
    "{{SPECIAL}}": "special",
    "{{THEME}}": "theme",
    "{{GROUP}}": "group",
    "{{NAME}}": "name",
    "{{NAME_TEACHER}}": "name_teacher",
    "{{NAME_NOR}}": "name_nor",
    "{{YEAR}}": "year"
}

# Global counters for sequential numbering
table_counter = 0
formula_counter = 0
image_counter = 0


def insert_section_break_after(paragraph):
    body = paragraph._p.getparent()
    section_properties = None

    for element in paragraph._p.itersiblings():
        sectPr = element.xpath(".//w:sectPr")
        if sectPr:
            section_properties = sectPr[0]
            break

    if section_properties is None:
        body_sectPr = body.xpath("./w:sectPr")
        if body_sectPr:
            section_properties = body_sectPr[0]

    if section_properties is None:
        return paragraph

    new_sectPr = deepcopy(section_properties)

    # тип разрыва = следующая страница
    type_el = new_sectPr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        new_sectPr.insert(0, type_el)

    type_el.set(qn("w:val"), "nextPage")

    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pPr.append(new_sectPr)
    p.append(pPr)

    paragraph._p.addnext(p)

    return Paragraph(p, paragraph._parent)


def set_list(paragraph, num_id, level=0):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")

    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))

    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))

    numPr.append(ilvl)
    numPr.append(numId)

    pPr.append(numPr)


def create_numbering(doc, base_num_id):
    """Creates a new w:num entry that references the same abstractNum as
    base_num_id but always restarts counting from 1 at level 0.
    This guarantees that every list block gets an independent counter."""
    numbering = doc.part.numbering_part.element
    base_num = numbering.xpath(f'./w:num[@w:numId="{base_num_id}"]')

    if not base_num:
        return base_num_id

    abstract_id = base_num[0].xpath("./w:abstractNumId/@w:val")[0]

    existing_ids = [
        int(num_id)
        for num_id in numbering.xpath("./w:num/@w:numId")
        if str(num_id).isdigit()
    ]
    new_num_id = max(existing_ids, default=0) + 1

    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))

    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    new_num.append(abstract)

    # Force restart at 1 for level 0 — critical for independent counters
    lvlOverride = OxmlElement("w:lvlOverride")
    lvlOverride.set(qn("w:ilvl"), "0")
    startOverride = OxmlElement("w:startOverride")
    startOverride.set(qn("w:val"), "1")
    lvlOverride.append(startOverride)
    new_num.append(lvlOverride)

    numbering.append(new_num)

    return new_num_id


def replace_xml_placeholders(template_path, output_path, data):
    replacements = {
        "{{SPECIAL}}": data["special"],
        "{{THEME}}": data["theme"],
        "{{GROUP}}": data["group"],
        "{{NAME}}": data["name"],
        "{{NAME_TEACHER}}": data["name_teacher"],
        "{{NAME_NOR}}": data["name_nor"],
        "{{YEAR}}": data["year"]
    }
    temp_dir = tempfile.mkdtemp(prefix="_docx_", dir=".")
    try:
        with ZipFile(template_path, "r") as zip_ref:
            zip_ref.extractall(temp_dir)

        for root, dirs, files in os.walk(temp_dir):
            for file in files:
                if not file.endswith(".xml"):
                    continue
                xml_path = os.path.join(root, file)
                try:
                    with open(xml_path, "r", encoding="utf-8") as f:
                        xml = f.read()
                    changed = False
                    for placeholder, value in replacements.items():
                        if placeholder in xml:
                            xml = xml.replace(
                                placeholder,
                                str(value)
                            )
                            changed = True
                    if changed:
                        with open(xml_path, "w", encoding="utf-8") as f:
                            f.write(xml)
                except:
                    pass
        with ZipFile(output_path, "w") as zip_out:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    full_path = os.path.join(root, file)
                    archive_name = os.path.relpath(
                        full_path,
                        temp_dir
                    ).replace("\\", "/")
                    zip_out.write(
                        full_path,
                        archive_name
                    )
    finally:
        shutil.rmtree(temp_dir)


def find_paragraph(doc, marker):
    for paragraph in doc.paragraphs:
        if marker in paragraph.text:
            return paragraph
    return None


def insert_after(paragraph, text="", style=None):
    parts = str(text).split("\n")
    current = paragraph
    for part in parts:
        new_p = OxmlElement("w:p")
        current._p.addnext(new_p)
        new_paragraph = Paragraph(new_p, paragraph._parent)
        new_paragraph.text = part
        if style:
            new_paragraph.style = style
        current = new_paragraph
    return current


def format_cell_paragraph(p, text, is_header=False):
    p.text = str(text)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    if is_header:
        p.alignment = 1  # Center alignment
    else:
        p.alignment = 0  # Left alignment

    # Set font family and size, explicitly disable bolding
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False


def set_formula_paragraph(p, alignment=0):
    """Apply formula-table paragraph formatting: all indents/spaces to 0,
    single line spacing, no first-line indent."""
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = alignment  # 0=left, 1=center, 2=right


def set_row_height_exact(row, height_cm):
    """Set exact row height (режим Точно) in cm."""
    from docx.shared import Cm as _Cm
    from docx.oxml.ns import qn as _qn
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    # 1 cm = 567 twips
    twips = int(height_cm * 567)
    trHeight.set(_qn('w:val'), str(twips))
    trHeight.set(_qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def set_table_cell_margins_zero(table):
    """Set all cell margins to 0 for a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), '0')
        el.set(qn('w:type'), 'dxa')
        tblCellMar.append(el)
    # Remove existing tblCellMar if present
    existing = tblPr.find(qn('w:tblCellMar'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblCellMar)


def remove_table_borders(table):
    """Remove all visible borders from a table (transparent/none)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Font family and size (Times New Roman 14pt)
    font_name = OxmlElement('w:rFonts')
    font_name.set(qn('w:ascii'), 'Times New Roman')
    font_name.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(font_name)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '28')  # 14pt (28 half-points)
    rPr.append(sz)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Color (standard blue for hyperlinks: 0000FF)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)

    new_run.append(rPr)

    new_text = OxmlElement('w:t')
    new_text.text = text
    new_run.append(new_text)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def format_list_item_text(text, is_last):
    t = str(text).strip()
    # Strip any trailing list punctuation (. ; ,)
    while t and t[-1] in ('.', ';', ','):
        t = t[:-1].strip()
    
    # Add trailing semicolon or period
    if is_last:
        return t + "."
    else:
        return t + ";"


def render_block(doc, anchor, block):
    global table_counter, formula_counter, image_counter
    block_type = block["type"]

    if block_type == "text":
        anchor = insert_after(
            anchor,
            block["text"],
            "Normal"
        )

    elif block_type == "bullet_list":
        num_id = create_numbering(doc, 6)
        items = block["items"]

        for idx, item in enumerate(items):
            formatted_text = format_list_item_text(item, idx == len(items) - 1)
            anchor = insert_after(
                anchor,
                formatted_text,
                "Normal"
            )
            set_list(anchor, num_id)

    elif block_type == "numbered_list":
        num_id = create_numbering(doc, 7)
        items = block["items"]

        for idx, item in enumerate(items):
            formatted_text = format_list_item_text(item, idx == len(items) - 1)
            anchor = insert_after(
                anchor,
                formatted_text,
                "Normal"
            )
            set_list(anchor, num_id)

    elif block_type == "table":
        table_counter += 1

        # 1. Create table caption/title
        table_title = block.get("title", "")
        caption_text = f"Таблица {table_counter} – {table_title}" if table_title else f"Таблица {table_counter}"

        # Before the table title there must be an enter (empty paragraph)
        caption_p_before = insert_after(anchor, "", style="Normal")
        caption_p = insert_after(caption_p_before, caption_text, style="Normal")

        # Table title: apply Normal style parameters explicitly
        # All indents = 0, spaces = 0, first_line_indent = 1 cm (red line / красная строка)
        caption_p.paragraph_format.first_line_indent = Cm(1.0)
        caption_p.paragraph_format.left_indent = Pt(0)
        caption_p.paragraph_format.right_indent = Pt(0)
        caption_p.paragraph_format.space_before = Pt(0)
        caption_p.paragraph_format.space_after = Pt(0)
        caption_p.paragraph_format.keep_with_next = True

        for run in caption_p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

        # 2. Create the table
        rows = block.get("rows", [])
        headers = block.get("headers", [])

        num_cols = len(headers) if headers else (len(rows[0]) if rows else 1)

        table = doc.add_table(rows=0, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Move table after the caption paragraph
        caption_p._p.addnext(table._tbl)

        # 3. Add headers if present
        if headers:
            row = table.add_row()
            for i, val in enumerate(headers):
                cell = row.cells[i]
                p = cell.paragraphs[0]
                format_cell_paragraph(p, val, is_header=True)

        # 4. Add data rows
        for r_val in rows:
            row = table.add_row()
            for i, val in enumerate(r_val):
                if i < len(row.cells):
                    cell = row.cells[i]
                    p = cell.paragraphs[0]
                    format_cell_paragraph(p, val, is_header=False)

        # 5. Create paragraph after the table for subsequent insertions
        after_p = OxmlElement("w:p")
        table._tbl.addnext(after_p)
        anchor = Paragraph(after_p, anchor._parent)
        anchor.paragraph_format.space_before = Pt(12)

    elif block_type == "formula":
        formula_counter += 1

        variables = block.get("variables", block.get("where", []))
        has_vars = bool(variables)

        # Formula text: append comma if variables follow
        formula_text = block.get("formula", block.get("text", "")).rstrip()
        if has_vars:
            while formula_text and formula_text[-1] in ('.', ';', ','):
                formula_text = formula_text[:-1].rstrip()
            formula_text += ","

        def _apply_font(p):
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

        # -------------------------------------------------------
        # Formula table: always 1 row, 2 cols [16cm formula | 1cm num]
        # -------------------------------------------------------
        ftbl = doc.add_table(rows=1, cols=2)
        ftbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_cell_margins_zero(ftbl)
        remove_table_borders(ftbl)

        ftbl.columns[0].width = Cm(16)
        ftbl.columns[1].width = Cm(1)
        ftbl.rows[0].cells[0].width = Cm(16)
        ftbl.rows[0].cells[1].width = Cm(1)
        set_row_height_exact(ftbl.rows[0], 0.6)

        pf = ftbl.rows[0].cells[0].paragraphs[0]
        pf.text = formula_text
        set_formula_paragraph(pf, alignment=1)  # center
        _apply_font(pf)

        pn = ftbl.rows[0].cells[1].paragraphs[0]
        pn.text = f"({formula_counter})"
        set_formula_paragraph(pn, alignment=1)  # center
        _apply_font(pn)

        anchor._p.addnext(ftbl._tbl)
        last_tbl = ftbl._tbl

        if has_vars:
            # -------------------------------------------------------
            # Variables table: N rows, 4 cols [1cm | 1cm | 0.5cm | 15cm]
            # Placed directly after formula table — visually one block
            # -------------------------------------------------------
            total_vars = len(variables)
            col_w = [Cm(1), Cm(1), Cm(0.5), Cm(15)]

            vtbl = doc.add_table(rows=total_vars, cols=4)
            vtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_cell_margins_zero(vtbl)
            remove_table_borders(vtbl)

            for i, w in enumerate(col_w):
                vtbl.columns[i].width = w

            for vi, item in enumerate(variables):
                row = vtbl.rows[vi]
                set_row_height_exact(row, 0.6)
                for i, w in enumerate(col_w):
                    row.cells[i].width = w

                if isinstance(item, dict):
                    symbol = item.get("symbol", item.get("name", ""))
                    description = item.get("description", item.get("text", ""))
                else:
                    symbol = ""
                    description = str(item)

                desc = description.rstrip()
                while desc and desc[-1] in ('.', ';', ','):
                    desc = desc[:-1].rstrip()
                desc += "." if vi == total_vars - 1 else ";"

                p0 = row.cells[0].paragraphs[0]
                p0.text = "где" if vi == 0 else ""
                set_formula_paragraph(p0, alignment=1)
                _apply_font(p0)

                p1 = row.cells[1].paragraphs[0]
                p1.text = symbol
                set_formula_paragraph(p1, alignment=1)
                _apply_font(p1)

                p2 = row.cells[2].paragraphs[0]
                p2.text = "–" if (symbol and description) else ""
                set_formula_paragraph(p2, alignment=0)
                _apply_font(p2)

                # col3: description — left
                p3 = row.cells[3].paragraphs[0]
                p3.text = desc
                set_formula_paragraph(p3, alignment=0)  # left
                _apply_font(p3)

            empty_p = OxmlElement("w:p")
            ftbl._tbl.addnext(empty_p)

            # превращаем в объект Paragraph
            empty_paragraph = Paragraph(empty_p, anchor._parent)

            # стиль Normal
            empty_paragraph.style = doc.styles['Normal']
            empty_p.addnext(vtbl._tbl)
            last_tbl = vtbl._tbl

        # Paragraph after the formula block for subsequent insertions
        after_p = OxmlElement("w:p")
        last_tbl.addnext(after_p)
        anchor = Paragraph(after_p, anchor._parent)

    elif block_type == "image":
        # 1. Сверху пустая строка Normal
        empty_p_before = insert_after(anchor, "", "Normal")

        # 2. Снизу фото (в новом абзаце)
        pic_p = insert_after(empty_p_before, "", "Normal")
        pic_p.alignment = 1  # Center alignment
        pic_p.paragraph_format.first_line_indent = Pt(0)
        pic_p.paragraph_format.left_indent = Pt(0)
        pic_p.paragraph_format.right_indent = Pt(0)
        pic_p.paragraph_format.space_before = Pt(0)
        pic_p.paragraph_format.space_after = Pt(0)

        img_path = block.get("src", "")
        if img_path:
            img_path = img_path.replace("\\", "/")
            if img_path.startswith("./"):
                img_path = img_path[2:]

        resolved_path = img_path
        if not (resolved_path and os.path.exists(resolved_path)):
            proj_dir = os.environ.get("PROJECT_DIR")
            if proj_dir:
                alt_path = os.path.join(proj_dir, img_path)
                if os.path.exists(alt_path):
                    resolved_path = alt_path

        if resolved_path and os.path.exists(resolved_path):
            try:
                run = pic_p.add_run()
                run.add_picture(resolved_path, width=Cm(15))
            except Exception as e:
                run = pic_p.add_run()
                run.text = f"[Ошибка загрузки изображения: {str(e)}]"
        else:
            run = pic_p.add_run()
            proj_dir = os.environ.get("PROJECT_DIR", "None")
            cwd = os.getcwd()
            run.text = f"[Изображение не найдено. Имя: {img_path}. CWD: {cwd}. PROJECT_DIR: {proj_dir}]"

        # 3. Пустая строка Normal
        empty_p_after = insert_after(pic_p, "", "Normal")

        # 4. Название Рисунок (счетчик) - (название)
        image_counter += 1
        caption = block.get("caption", "").strip()
        
        import re
        caption_clean = re.sub(r'^Рисунок\s*\d*[\s\-\—\–]*', '', caption, flags=re.IGNORECASE).strip()
        caption_text = f"Рисунок {image_counter} – {caption_clean}" if caption_clean else f"Рисунок {image_counter}"

        caption_p = insert_after(empty_p_after, caption_text, "Normal")
        caption_p.alignment = 1  # Center alignment
        caption_p.paragraph_format.first_line_indent = Pt(0)
        caption_p.paragraph_format.left_indent = Pt(0)
        caption_p.paragraph_format.right_indent = Pt(0)
        caption_p.paragraph_format.space_before = Pt(0)
        caption_p.paragraph_format.space_after = Pt(0)

        for run in caption_p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

        anchor = caption_p

    return anchor


if __name__ == '__main__':
    with open("data.json", "r", encoding="utf-8") as f:
        data = json.load(f)

    # Определяем имя файла шаблона по значению из data.json
    template_type = data.get("template", "D")  # по умолчанию D
    template_file = f"template{template_type}.docx"
    
    # На всякий случай проверим, что файл существует, иначе откат на templateD.docx
    if not os.path.exists(template_file):
        template_file = "templateD.docx"

    replace_xml_placeholders(
        template_file,
        "temp.docx",
        data
    )

    doc = Document("temp.docx")

    # -------------------
    # ВВЕДЕНИЕ
    # -------------------
    intro_marker = find_paragraph(doc, "{{INTRODUCTION}}")
    if intro_marker:
        introduction = data["introduction"]
        if isinstance(introduction, list):
            # New format: array of content blocks (same as section content)
            current = intro_marker
            current.text = ""
            for block in introduction:
                current = render_block(doc, current, block)
            # Remove the empty placeholder paragraph
            p = intro_marker._p
            p.getparent().remove(p)
        else:
            # Legacy format: plain string with \n line breaks
            parts = str(introduction).split("\n")
            intro_marker.text = parts[0]
            current = intro_marker
            for part in parts[1:]:
                current = insert_after(current, part, "Normal")

    # -------------------
    # РАЗДЕЛЫ
    # -------------------
    sections_marker = find_paragraph(doc, "{{SECTIONS}}")
    if sections_marker:
        current = sections_marker

        for section in data["sections"]:
            current = insert_after(
                current,
                section["title"]["text"],
                "Heading 1"
            )
            # Add an empty paragraph after Heading 1
            current = insert_after(current, "", "Normal")

            for block in section["content"]:
                current = render_block(doc, current, block)

            for subsection in section.get("subsections", []):
                current = insert_after(
                    current,
                    subsection["title"]["text"],
                    "Heading 2"
                )

                for block in subsection["content"]:
                    current = render_block(doc, current, block)

            current = insert_section_break_after(current)

        # Delete sections_marker to avoid extra Enter
        p = sections_marker._p
        p.getparent().remove(p)

    # -------------------
    # ЗАКЛЮЧЕНИЕ
    # -------------------
    conclusion_marker = find_paragraph(doc, "{{CONCLUSION}}")
    if conclusion_marker:
        conclusion = data.get("conclusion", [])
        if isinstance(conclusion, list):
            current = conclusion_marker
            current.text = ""
            for block in conclusion:
                current = render_block(doc, current, block)
            # Delete conclusion_marker to avoid extra Enter
            p = conclusion_marker._p
            p.getparent().remove(p)
        else:
            # Legacy format: plain string with \n line breaks
            parts = str(conclusion).split("\n")
            conclusion_marker.text = parts[0]
            current = conclusion_marker
            for part in parts[1:]:
                current = insert_after(current, part, "Normal")

    # -------------------
    # ЛИТЕРАТУРА
    # -------------------
    reference_marker = find_paragraph(doc, "{{REFERENCE}}")
    if reference_marker:
        current = reference_marker
        num_id = create_numbering(doc, 5)
        ref_items = data["reference"]

        for idx, item in enumerate(ref_items):
            is_last = (idx == len(ref_items) - 1)
            # Middle dash (en-dash \u2013) between title and URL
            title_part = f'{item["text"]} \u2013 '
            current = insert_after(
                current,
                title_part,
                "Normal"
            )
            set_list(current, num_id)
            # Append the URL as a hyperlink
            url_display = item["url"] + ("." if is_last else ";")
            add_hyperlink(current, url_display, item["url"])

        # Delete reference_marker to avoid extra Enter
        p = reference_marker._p
        p.getparent().remove(p)

    # -------------------
    # СОХРАНЕНИЕ
    # -------------------
    doc.save("result.docx")
    os.remove("temp.docx")
    print("Готово!")
